"""Tests for the Update Email Generator (build_email_doc)."""

from pathlib import Path

import pytest

from app.email_export import EmailData, OtherItem, ScopeSection, build_email_doc


def _minimal_data(**overrides) -> EmailData:
    base = dict(
        boiler_name="RECOVERY BOILER #2",
        status_date="6/14/2026",
        status_time="7:30 PM",
        overall_summary="Total initial data turned over to BSI is 0%.",
        scope_work_findings=[],
        visual_findings=[],
        scope_sections=[
            ScopeSection(name="Floor UT", status="No initial data received."),
        ],
        other_scope_items=[],
        punchlist_items=[],
    )
    base.update(overrides)
    return EmailData(**base)


def test_build_email_doc_creates_file(tmp_path):
    data = _minimal_data()
    out = build_email_doc(data, tmp_path / "email.docx")
    assert out.exists()
    assert out.stat().st_size > 0


def test_build_email_doc_contains_boiler_name(tmp_path):
    from docx import Document
    data = _minimal_data()
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "RECOVERY BOILER #2" in texts


def test_build_email_doc_legend_paragraphs(tmp_path):
    from docx import Document
    data = _minimal_data()
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "COMPLETE" in texts
    assert "IN PROGRESS" in texts
    assert "ISSUES NOTED" in texts
    assert "OTHER ISSUES" in texts


def test_build_email_doc_scope_section_line(tmp_path):
    from docx import Document
    data = _minimal_data(scope_sections=[
        ScopeSection(name="Floor UT", status="No initial data received."),
        ScopeSection(name="Front Wall MLO", status="In progress."),
    ])
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Floor UT – No initial data received." in texts
    assert "Front Wall MLO – In progress." in texts


def test_build_email_doc_findings_empty_shows_none(tmp_path):
    from docx import Document
    data = _minimal_data(scope_work_findings=[], visual_findings=[])
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert texts.count("None reported at this time.") == 2


def test_build_email_doc_findings_shown(tmp_path):
    from docx import Document
    data = _minimal_data(scope_work_findings=["Economizer tube 79 cut out."])
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Economizer tube 79 cut out." in texts
    assert texts.count("None reported at this time.") == 1


def test_build_email_doc_aux_items_shown_when_present(tmp_path):
    from docx import Document
    data = _minimal_data(
        other_scope_items=[OtherItem(description="PT OF COMPOSITE PORTS", status="complete")]
    )
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Other scope items:" in texts
    assert "PT OF COMPOSITE PORTS – complete" in texts


def test_build_email_doc_other_scope_empty_shows_placeholder(tmp_path):
    from docx import Document
    data = _minimal_data(other_scope_items=[])
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Other scope items:" in texts
    assert "No other scope items." in texts


def test_build_email_doc_punchlist_shown_when_present(tmp_path):
    from docx import Document
    data = _minimal_data(
        punchlist_items=[OtherItem(description="Item 37 – UT spout 2 tube 38", status="complete")]
    )
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Punchlist:" in texts
    assert "Item 37 – UT spout 2 tube 38 – complete" in texts


def test_build_email_doc_punchlist_empty_shows_placeholder(tmp_path):
    from docx import Document
    data = _minimal_data(punchlist_items=[])
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Punchlist:" in texts
    assert "No punchlist items." in texts


def test_build_email_doc_other_issues_legend_purple(tmp_path):
    from docx import Document
    from docx.oxml.ns import qn
    data = _minimal_data()
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "OTHER ISSUES" in texts
    p = next(para for para in doc.paragraphs if para.text == "OTHER ISSUES")
    fills = []
    for r in p.runs:
        shd = r._r.find(f".//{{{r._r.nsmap.get('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')}}}shd")
        if shd is not None:
            fills.append(shd.get(qn("w:fill")))
    assert "CC99FF" in fills


def test_build_email_doc_page_size_letter(tmp_path):
    from docx import Document
    from docx.shared import Inches
    data = _minimal_data()
    out = build_email_doc(data, tmp_path / "email.docx")
    doc = Document(str(out))
    section = doc.sections[0]
    assert abs(section.page_width.inches - 8.5) < 0.01
    assert abs(section.page_height.inches - 11.0) < 0.01
