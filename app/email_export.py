"""Build formatted .docx update email documents using python-docx."""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


@dataclass
class ScopeSection:
    """One boiler section and its current status line for the email."""

    name: str
    status: str = "No initial data received."


@dataclass
class OtherItem:
    """One auxiliary scope or punchlist item and its current status."""

    description: str
    status: str = "no report received"


@dataclass
class EmailData:
    """All data needed to render an NDE status update email."""

    boiler_name: str
    status_date: str
    status_time: str
    overall_summary: str
    scope_work_findings: list[str] = field(default_factory=list)
    visual_findings: list[str] = field(default_factory=list)
    scope_sections: list[ScopeSection] = field(default_factory=list)
    other_scope_items: list[OtherItem] = field(default_factory=list)
    punchlist_items: list[OtherItem] = field(default_factory=list)


# --- Date formatting ----------------------------------------------------------

def _format_date(date_str: str) -> str:
    """Convert ISO ('2026-06-15') or M/D/YYYY ('6/15/2026') to 'Month D, YYYY'."""
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d")
        return f"{d.strftime('%B')} {d.day}, {d.year}"
    except ValueError:
        pass
    try:
        parts = date_str.split("/")
        if len(parts) == 3:
            d = datetime(int(parts[2]), int(parts[0]), int(parts[1]))
            return f"{d.strftime('%B')} {d.day}, {d.year}"
    except (ValueError, IndexError):
        pass
    return date_str


# --- Separator ----------------------------------------------------------------

_SEP = " – "  # en-dash with spaces


# --- XML helpers for run shading ----------------------------------------------

def _shade_run(run, hex_color: str) -> None:
    """Apply character background shading to a run via XML."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


# --- Document construction ----------------------------------------------------

def _add_blank(doc: Document) -> None:
    doc.add_paragraph("")


def _add_bold(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def _add_shaded_legend(doc: Document, text: str, fill: str | None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if fill:
        _shade_run(run, fill)


def _add_findings(doc: Document, findings: list[str]) -> None:
    if not findings:
        doc.add_paragraph("None reported at this time.", style="List Paragraph")
    else:
        for finding in findings:
            doc.add_paragraph(finding, style="List Paragraph")


def _add_other_items(doc: Document, items: list[OtherItem]) -> None:
    for item in items:
        doc.add_paragraph(item.description + _SEP + item.status, style="List Paragraph")


def build_email_doc(data: EmailData, output_path: str | Path) -> Path:
    """Generate a formatted update email .docx and save to output_path."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Page setup: US Letter, 1" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default font and paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    list_style = doc.styles["List Paragraph"]
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(0)
    list_style.paragraph_format.line_spacing = 1.0

    # Opening
    doc.add_paragraph("All,")
    _add_blank(doc)
    doc.add_paragraph(
        f"Please see below for status of NDT inspection as of "
        f"{data.status_time} on {_format_date(data.status_date)}. "
        f"If you have any questions, please don't hesitate to reach out."
    )
    _add_blank(doc)

    # Legend
    _add_shaded_legend(doc, "COMPLETE", "00FF00")
    _add_shaded_legend(doc, "IN PROGRESS", "FFFF00")
    _add_shaded_legend(doc, "ISSUES NOTED", "FF0000")
    _add_shaded_legend(doc, "Other Issues", None)
    _add_blank(doc)

    # Boiler heading
    _add_bold(doc, data.boiler_name)
    _add_blank(doc)

    # Overall summary
    doc.add_paragraph(data.overall_summary)
    _add_blank(doc)

    # Discovery — scope work
    _add_bold(doc, "Discovery based on scope work:")
    _add_findings(doc, data.scope_work_findings)
    _add_blank(doc)

    # Discovery — visual
    _add_bold(doc, "Discovery based on visual inspection:")
    _add_findings(doc, data.visual_findings)
    _add_blank(doc)

    # Scope of work
    _add_bold(doc, "Scope of work:")
    for s in data.scope_sections:
        doc.add_paragraph(s.name + _SEP + s.status)
    _add_blank(doc)

    # Other scope items
    if data.other_scope_items:
        _add_bold(doc, "Other scope items:")
        _add_blank(doc)
        _add_other_items(doc, data.other_scope_items)
        _add_blank(doc)

    # Punchlist
    if data.punchlist_items:
        _add_bold(doc, "Punchlist")
        _add_blank(doc)
        for item in data.punchlist_items:
            doc.add_paragraph(item.description + _SEP + item.status)

    doc.save(str(output_path))
    return output_path
